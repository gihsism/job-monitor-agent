"""Cover letter generation module using Claude API."""

import json
import re
from datetime import datetime
from pathlib import Path

import anthropic
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import MASTER_CV

OUTPUT_DIR = Path(__file__).parent / "output"

COVER_LETTER_PROMPT = """\
You are an expert cover letter writer. Write a tailored cover letter for Alena Nikolskaia for the job below.

## Alena's Background
{master_cv}

## Job Posting
Title: {job_title}
Company: {company}
URL: {job_url}

Description:
{job_description}

## Instructions

Write a professional, concise cover letter (300-400 words) following these rules:

1. **DO NOT INVENT anything.** Only use facts from Alena's background above.
2. Open with a strong hook connecting her background to this specific role.
3. Highlight her unique value: finance expertise + AI/tech skills (ETH CAS in AI, AI platform at KPMG).
4. Show concrete achievements with numbers where possible.
5. Close with enthusiasm and a call to action.
6. Tone: confident, professional, not generic. Avoid clichés.

Return a JSON object with:
- "greeting": string (e.g., "Dear Hiring Manager,")
- "body": list of strings (each string is a paragraph)
- "closing": string (e.g., "Kind regards,")
- "company_name": string (extracted from the job posting)

Return ONLY valid JSON, no markdown fences.
"""


def _extract_company(job: dict) -> str:
    """Try to extract company name from job title or URL."""
    url = job.get("url", "")
    title = job.get("title", "")

    # Common patterns: "Role at Company", "Role - Company"
    for sep in [" at ", " - ", " | ", " — "]:
        if sep in title:
            return title.split(sep)[-1].strip()

    # Try domain
    for domain in ["linkedin.com", "indeed.ch", "jobs.ch", "glassdoor.com"]:
        if domain in url:
            return "the company"

    # Use domain name as fallback
    from urllib.parse import urlparse
    try:
        host = urlparse(url).hostname or ""
        parts = host.replace("www.", "").split(".")
        if parts and parts[0] not in ("jobs", "boards", "careers"):
            return parts[0].capitalize()
    except Exception:
        pass

    return "the company"


def generate_cover_letter(anthropic_api_key: str, job: dict) -> dict | None:
    """Generate a tailored cover letter. Returns cover letter dict."""
    client = anthropic.Anthropic(api_key=anthropic_api_key)
    company = _extract_company(job)

    prompt = COVER_LETTER_PROMPT.format(
        master_cv=json.dumps(MASTER_CV, indent=2),
        job_title=job["title"],
        company=company,
        job_url=job["url"],
        job_description=job["text"][:4000],
    )

    try:
        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}],
        )

        text = response.content[0].text.strip()
        if text.startswith("```"):
            text = text.split("\n", 1)[1]
            if text.endswith("```"):
                text = text[:-3]

        return json.loads(text)

    except Exception as e:
        print(f"[ERROR] Cover letter generation failed: {e}")
        return None


def generate_cover_letter_docx(cover_letter: dict, job_title: str) -> str:
    """Generate a .docx cover letter. Returns the file path."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("Alena Nikolskaia")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)

    contact = doc.add_paragraph()
    run = contact.add_run("+41 (0) 79 375 13 30  |  linkedin.com/in/alena-n-80966153")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Date
    date_p = doc.add_paragraph()
    run = date_p.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.size = Pt(10)

    doc.add_paragraph()  # spacing

    # Greeting
    greeting = doc.add_paragraph()
    run = greeting.add_run(cover_letter.get("greeting", "Dear Hiring Manager,"))
    run.font.size = Pt(11)

    # Body paragraphs
    for para_text in cover_letter.get("body", []):
        p = doc.add_paragraph(para_text)
        for run in p.runs:
            run.font.size = Pt(11)

    # Closing
    doc.add_paragraph()
    closing = doc.add_paragraph()
    run = closing.add_run(cover_letter.get("closing", "Kind regards,"))
    run.font.size = Pt(11)

    name = doc.add_paragraph()
    run = name.add_run("Alena Nikolskaia")
    run.bold = True
    run.font.size = Pt(11)

    # Save
    OUTPUT_DIR.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = re.sub(r"[^\w\s-]", "", job_title)[:50].strip().replace(" ", "_")
    filename = f"CoverLetter_Nikolskaia_{safe_title}_{timestamp}.docx"
    filepath = OUTPUT_DIR / filename

    doc.save(str(filepath))
    return str(filepath)
