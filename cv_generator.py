"""Generate a Word document CV from tailored CV data."""

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = Path(__file__).parent / "output"


def _add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
    return heading


def _add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def generate_cv_docx(tailored_cv: dict, job_title: str, job_url: str) -> str:
    """Generate a .docx CV file. Returns the file path."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # === Header ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ALENA NIKOLSKAIA")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("ACCA, Swiss Audit License")
    run.font.size = Pt(11)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run("+41 (0) 79 375 13 30  |  linkedin.com/in/alena-n-80966153")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # === Tailored for ===
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(f"CV tailored for: {job_title}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    # === Summary ===
    _add_heading(doc, "Summary")
    p = doc.add_paragraph(tailored_cv.get("tailored_summary", ""))
    for run in p.runs:
        run.font.size = Pt(10)

    # === Key Skills ===
    skills = tailored_cv.get("key_skills", [])
    if skills:
        _add_heading(doc, "Key Skills")
        # Show as comma-separated for compactness
        p = doc.add_paragraph(" • ".join(skills))
        for run in p.runs:
            run.font.size = Pt(10)

    # === Professional Experience ===
    _add_heading(doc, "Professional Experience")
    for exp in tailored_cv.get("experience", []):
        # Company header
        p = doc.add_paragraph()
        run = p.add_run(f"{exp.get('company', '')} ({exp.get('location', '')})")
        run.bold = True
        run.font.size = Pt(11)
        # Period on same line
        run = p.add_run(f"    {exp.get('period', '')}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Title
        p = doc.add_paragraph()
        run = p.add_run(exp.get("title", ""))
        run.italic = True
        run.font.size = Pt(10)

        # Bullets
        for bullet in exp.get("bullets", []):
            _add_bullet(doc, bullet)

    # === Education ===
    _add_heading(doc, "Education & Certifications")
    for edu in tailored_cv.get("education", []):
        _add_bullet(doc, edu)
    for cert in tailored_cv.get("certifications", []):
        _add_bullet(doc, cert)

    # === Languages ===
    _add_heading(doc, "Languages")
    langs = tailored_cv.get("languages", [])
    p = doc.add_paragraph(" | ".join(langs))
    for run in p.runs:
        run.font.size = Pt(10)

    # Save
    OUTPUT_DIR.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = re.sub(r"[^\w\s-]", "", job_title)[:50].strip().replace(" ", "_")
    filename = f"CV_Nikolskaia_{safe_title}_{timestamp}.docx"
    filepath = OUTPUT_DIR / filename

    doc.save(str(filepath))
    return str(filepath)
